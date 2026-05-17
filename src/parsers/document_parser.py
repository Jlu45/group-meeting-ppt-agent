import os
import logging
from pathlib import Path
from typing import Optional

from src.models import ParsedDocument

logger = logging.getLogger(__name__)

PDF_SIZE_THRESHOLD = 5 * 1024 * 1024
TABLE_MARKER_THRESHOLD = 5


class UniversalDocumentParser:
    """通用文档解析器 - 任意格式 → Markdown"""

    def __init__(self, use_docling: bool = False):
        self._use_docling = use_docling

    def parse(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        file_type = ext.lstrip(".")

        if ext == ".pdf":
            content = self._parse_pdf(file_path)
        elif ext == ".docx":
            content = self._parse_docx(file_path)
        elif ext == ".xlsx":
            content = self._parse_xlsx(file_path)
        elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff"):
            content = self._parse_image(file_path)
        elif ext in (".md", ".markdown"):
            content = path.read_text(encoding="utf-8")
        elif ext in (".txt", ".csv", ".json", ".yaml", ".yml"):
            content = path.read_text(encoding="utf-8")
        elif ext in (".html", ".htm"):
            content = self._parse_html(file_path)
        else:
            content = self._parse_text_fallback(file_path)

        logger.info(f"Parsed {file_path} -> {len(content)} chars")
        return ParsedDocument(
            source_path=str(path.resolve()),
            markdown_content=content,
            file_type=file_type,
        )

    def parse_multiple(self, file_paths: list[str]) -> list[ParsedDocument]:
        results = []
        for fp in file_paths:
            try:
                results.append(self.parse(fp))
            except Exception as e:
                logger.error(f"Failed to parse {fp}: {e}")
        return results

    def merge_documents(self, documents: list[ParsedDocument]) -> ParsedDocument:
        if not documents:
            return ParsedDocument(source_path="", markdown_content="")

        merged_content = "\n\n---\n\n".join(
            f"## Source: {Path(doc.source_path).name}\n\n{doc.markdown_content}"
            for doc in documents
        )
        return ParsedDocument(
            source_path=";".join(d.source_path for d in documents),
            markdown_content=merged_content,
            file_type="merged",
        )

    def _parse_pdf(self, path: str) -> str:
        try:
            import PyPDF2
            content = ""
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        content += text + "\n\n"
            return content.strip()
        except Exception as e:
            logger.warning(f"PyPDF2 failed, trying pdfplumber: {e}")
            try:
                import pdfplumber
                with pdfplumber.open(path) as pdf:
                    content = ""
                    for page in pdf.pages:
                        content += page.extract_text() + "\n\n"
                return content.strip()
            except Exception as e2:
                logger.error(f"Failed to parse PDF {path}: {e2}")
                return f"[PDF File: {Path(path).name}]"

    def _parse_docx(self, path: str) -> str:
        try:
            from docx import Document
            doc = Document(path)
            content = "\n\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
            return content
        except Exception as e:
            logger.error(f"Failed to parse DOCX {path}: {e}")
            return f"[DOCX File: {Path(path).name}]"

    def _parse_xlsx(self, path: str) -> str:
        try:
            import pandas as pd
            df = pd.read_excel(path)
            return df.to_markdown(index=False)
        except Exception as e:
            logger.error(f"Failed to parse XLSX {path}: {e}")
            return f"[XLSX File: {Path(path).name}]"

    def _parse_image(self, path: str) -> str:
        try:
            from PIL import Image
            img = Image.open(path)
            return f"![{Path(path).stem}]({path})\n\nImage: {Path(path).name}, Size: {img.size[0]}x{img.size[1]}"
        except Exception as e:
            return f"[Image: {Path(path).name}] (Error: {e})"

    def _parse_html(self, path: str) -> str:
        try:
            from bs4 import BeautifulSoup
            with open(path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f, "html.parser")
                return soup.get_text()
        except Exception as e:
            logger.warning(f"Failed to parse HTML {path}: {e}")
            try:
                return path.read_text(encoding="utf-8")
            except Exception:
                return f"[HTML File: {Path(path).name}]"

    def _parse_text_fallback(self, path: str) -> str:
        try:
            return path.read_text(encoding="utf-8")
        except Exception as e:
            logger.error(f"Failed to parse {path}: {e}")
            return f"[File: {Path(path).name}]"
